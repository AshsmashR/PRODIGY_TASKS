# -*- coding: utf-8 -*-
"""TEXT GENERATION RNN

Automatically generated by Colab.

Original file is located at
    https://colab.research.google.com/drive/1Bu6SMwaT6uGoCBg0OtAEPCZ8xrjxxBhx
"""

!pip install python-docx
#we have to install this to load large dox!!

from docx import Document
import re
from tensorflow.keras.preprocessing.text import Tokenizer #type:ignore
from sklearn.model_selection import train_test_split
import numpy as np
from tensorflow.keras.utils import to_categorical #type:ignore

#lets install the neccessary libraries and load the dataset
#I have downloaded my favourite classic book Pride and Prejudice by Jane austen

doc = Document(r'/content/The Project Gutenberg eBook of Pride and Prejudice.docx')
text = []

for paragraph in doc.paragraphs:
    if paragraph.text.strip():
        text.append(paragraph.text)
text_data = ' '.join(text)  #lets join the text in a sentence format .join!!

#so i imported re to clean my text of spaces, any special chars and lowercased it
text_data = re.sub(r'([a-z])([A-Z])', r'\1 \2', text_data)
text_data = re.sub(r'\.{2,}([!?])', r'\1', text_data)
text_data = re.sub(r'\.{2,}', '.', text_data)
text_data = re.sub(r'\s+', ' ', text_data).strip()

#MY download has lot of error and i had to write so common errors to do .replace
#by the corresponding corrections
common_replacements = {
    "lovingby": "loving by",
    "appliesto": "applies to",
    "verynumerous": "very numerous",
    "itbrings": "it brings",
    "asto": "as to",
    "byallowance": "by allowance",
    "andproper": "and proper",
    "theright": "the right",
    "ove.ove.ove": ".",
}
for error, correction in common_replacements.items():
    text_data = text_data.replace(error, correction)

#tokenize the words ease for our RNN
# words into token = numbers assigned for each word
tkr = Tokenizer(char_level=False)
tkr.fit_on_texts([text_data])
sequences = tkr.texts_to_sequences([text_data])[0]
word_index = tkr.word_index
print(f"Vocabulary Size: {len(word_index)}")
print(f"First 20 Tokens: {sequences[:20]}")


#Vocabulary Size: 7247

sequence_length = 15 #for my sentence, 30-50 takes lot of computation power
input_sequences = []
output_words = []


#A for loop for adding input sq into output words
for i in range(len(sequences) - sequence_length):
    input_seq = sequences[i:i + sequence_length]
    output_word = sequences[i + sequence_length]
    input_sequences.append(input_seq)
    output_words.append(output_word)

X = np.array(input_sequences)
y = np.array(output_words)

from tensorflow.keras.utils import Sequence  #type: ignore
from sklearn.model_selection import train_test_split
import numpy as np




X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
print(f"Total Sequences: {len(input_sequences)}")
print(f"Training Data Shape (X_train): {X_train.shape}")
print(f"Testing Data Shape (X_test): {X_test.shape}")
print(f"Training Labels Shape (y_train): {y_train.shape}")
print(f"Testing Labels Shape (y_test): {y_test.shape}")


y_train = np.array(output_words[:len(X_train)])
y_test = np.array(output_words[len(X_train):])


print(f"Training Labels Shape (y_train): {y_train.shape}")
print(f"Testing Labels Shape (y_test): {y_test.shape}")

from tensorflow.keras.models import Sequential              #type: ignore
from tensorflow.keras.layers import Embedding, LSTM, Dense  #type: ignore


vocab_size = len(set(sequences)) + 1 # Adding 1 to account for padding (if any)
sequence_length = 15
model_1 = Sequential([
    Embedding(input_dim=vocab_size, output_dim=100),
    LSTM(128, return_sequences=False),
    Dense(vocab_size, activation='softmax')
])

model_1.build(input_shape=(None, sequence_length))

# Compile the model
model_1.compile(optimizer='adam', loss='sparse_categorical_crossentropy', metrics=['accuracy'])

# Print model summary
model_1.summary()

print("X_train shape:", X_train.shape)
print("y_train shape:", y_train.shape)

# @title 1st model
#Training my model with 10 epochs at first
history = model_1.fit(
    X_train,
    y_train,
    validation_data=(X_test, y_test),
    epochs=10,
    batch_size=64,
    verbose=1
)

#I did not save the model accuracy is really poor

# @title trying 2nd model
import numpy as np
import re
from tensorflow.keras.preprocessing.text import Tokenizer
from tensorflow.keras.preprocessing.sequence import pad_sequences
from tensorflow.keras.utils import to_categorical
from tensorflow.keras.models import Sequential
from tensorflow.keras.layers import Embedding, LSTM, Dense, Dropout
from tensorflow.keras.callbacks import EarlyStopping, ModelCheckpoint
from sklearn.model_selection import train_test_split

def preprocess_text(text):
    text = re.sub(r'[^a-zA-Z\s]', '', text)
    text = text.lower()
    text = re.sub(r'\s+', ' ', text).strip()
    return text
from docx import Document
doc = Document(r"/content/The Project Gutenberg eBook of Pride and Prejudice.docx")
text = []
for paragraph in doc.paragraphs:
    if paragraph.text.strip():
        text.append(paragraph.text)
text_data = ' '.join(text)
text_data = preprocess_text(text_data)

tokenizer = Tokenizer()
tokenizer.fit_on_texts([text_data])
sequences = tokenizer.texts_to_sequences([text_data])[0]
vocab_size = len(tokenizer.word_index) + 1
print(f"Vocabulary Size: {vocab_size}")

sequence_length = 30  # Increased context length
input_sequences = []
output_words = []

for i in range(len(sequences) - sequence_length):
    input_seq = sequences[i:i + sequence_length]
    output_word = sequences[i + sequence_length]
    input_sequences.append(input_seq)
    output_words.append(output_word)
X = np.array(input_sequences)
y = np.array(output_words)

X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
print(f"Training Shape: {X_train.shape}, Testing Shape: {X_test.shape}")

#added glove embeddings for depth of the labels in text generation

def load_glove_embeddings(glove_path, word_index, embedding_dim=100):
    embeddings_index = {}
    with open(glove_path, 'r', encoding='utf-8') as f:
        for line in f:
            values = line.split()
            word = values[0]
            coefficients = np.asarray(values[1:], dtype='float32')
            embeddings_index[word] = coefficients
    embedding_matrix = np.zeros((vocab_size, embedding_dim))
    for word, i in word_index.items():
        embedding_vector = embeddings_index.get(word)
        if embedding_vector is not None:
            embedding_matrix[i] = embedding_vector
    return embedding_matrix

# Path to GloVe file
glove_path = r"/content/glove.6B.50d.txt"
vocab_size = len(tokenizer.word_index) + 1  # Vocabulary size

embedding_dim = 50
embedding_matrix = load_glove_embeddings(glove_path, tokenizer.word_index, embedding_dim)
print("Embedding matrix shape:", embedding_matrix.shape)

print("Tokenizer vocab size:", len(tokenizer.word_index))

embedding_matrix

from tensorflow.keras.models import Sequential
from tensorflow.keras.layers import Embedding, LSTM, Dense

# Define parameters
vocab_size = len(set(sequences)) + 1  # Vocabulary size
sequence_length = 20  # Length of input sequences
embedding_dim = 50  # Dimensions of embeddings
model_2 = Sequential()


#with glove embeddings as weights
model_2.add(Embedding(
    input_dim=vocab_size,          # Vocabulary size
    output_dim=embedding_dim,      # Embedding dimensions
    weights=[embedding_matrix],    # Pre-trained embeddings
    input_length=sequence_length,  # Input sequence length
    trainable=False                # Keep embeddings fixed
))

model_2.add(LSTM(256, return_sequences=False))
model_2.add(Dense(vocab_size, activation='softmax'))
model_2.build(input_shape=(None, sequence_length))

model_2.compile(optimizer='adam',
                loss='sparse_categorical_crossentropy',
                metrics=['accuracy'])

# Display model summary
model_2.summary()

# Callbacks
#early_stopping = EarlyStopping(monitor='val_loss', patience=5, restore_best_weights=True)
#model_checkpoint = ModelCheckpoint("best_text_gen_model.keras", save_best_only=True, monitor='val_loss', mode='min')

# Train the model
history = model_2.fit(
    X_train,
    y_train,
    validation_data=(X_test, y_test),
    epochs=20,
    batch_size=64,
    #callbacks=[early_stopping, model_checkpoint],
    #verbose=1
)

#the behavior of increasing training accuracy and decreasing validation accuracy in an RNN for text generation can be attributed to several factors, a
#is likely cause is the nature of the dataset.

def generate_text(seed_text, next_words, model, tokenizer, sequence_length):
    for _ in range(next_words):
        token_list = tokenizer.texts_to_sequences([seed_text])[0]
        token_list = pad_sequences([token_list], maxlen=sequence_length, padding='pre')
        predicted = model.predict(token_list, verbose=0)
        predicted_word_index = np.argmax(predicted)
        output_word = tokenizer.index_word.get(predicted_word_index, '')
        seed_text += " " + output_word
    return seed_text

# Generate a sentence
seed_text = "what is the weather like"

generated_text = generate_text(seed_text, next_words=20, model=model_2, tokenizer=tokenizer, sequence_length=sequence_length)
print(f"Generated Text: {generated_text}")

#because we had no labels in the dataset used it is hard to get high accuracy, the model is successful in generating text but not meaningful text

from tensorflow.keras.utils import to_categorical #lets hot encode

y_train = to_categorical(y_train, num_classes=vocab_size)
y_test = to_categorical(y_test, num_classes=vocab_size)

from tensorflow.keras.models import Sequential
from tensorflow.keras.layers import Embedding, LSTM, Dense, Dropout, BatchNormalization, Input
from tensorflow.keras.optimizers import Adam
from tensorflow.keras.regularizers import l2

vocab_size = len(set(sequences)) + 1                         # Adding 1 to account for padding (if any)
sequence_length = 20
embedding_dim = 50

# Updated model with fine-tuning
model_3 = Sequential([
    Input(shape=(sequence_length,)),
    Embedding(
        input_dim=vocab_size,
        output_dim=embedding_dim,
        weights=[embedding_matrix],
        trainable=True
    ),
    LSTM(256, return_sequences=True, kernel_regularizer=l2(0.01)),
    Dropout(0.2),
    BatchNormalization(),
    LSTM(128, return_sequences=False),
    Dense(128, activation='relu', kernel_regularizer=l2(0.01)),
    Dense(vocab_size, activation='softmax')
])
optimizer = Adam(learning_rate=1e-3)
model_3.compile(optimizer=optimizer, loss='categorical_crossentropy', metrics=['accuracy'])
model_3.summary()

# Train the model
history = model_3.fit(
    X_train,
    y_train,
    validation_data=(X_test, y_test),
    epochs=20,
    batch_size=128,
    #callbacks=[early_stopping, model_checkpoint],
    verbose=1
)

#Generate text
def generate_text(seed_text, next_words, model, tokenizer, sequence_length):
    for _ in range(next_words):
        token_list = tokenizer.texts_to_sequences([seed_text])[0]
        token_list = pad_sequences([token_list], maxlen=sequence_length, padding='pre')
        predicted = model.predict(token_list, verbose=0)
        predicted_word_index = np.argmax(predicted)
        output_word = tokenizer.index_word.get(predicted_word_index, '')
        seed_text += " " + output_word
    return seed_text

# Generate a sentence
seed_text = "what is the weather like"
generated_text = generate_text(seed_text, next_words=20, model=model_3, tokenizer=tokenizer, sequence_length=sequence_length)
print(f"Generated Text: {generated_text}")

#model_3 is notbetter than 2